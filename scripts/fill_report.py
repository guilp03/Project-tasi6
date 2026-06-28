"""Preenche o template oficial do Relatório Final IF1015.

O script clona o template .docx, substitui os placeholders da capa e das
seções, preenche as tabelas de economicidade e gera o arquivo final
`Relatorio_Final_Equipe3.docx`.
"""

import shutil
import warnings
from pathlib import Path

from docx import Document

# python-docx emite warning interno ao resolver estilos por style_id em
# alguns templates; o comportamento é correto, então suprimimos.
warnings.filterwarnings("ignore", message="style lookup by style_id")
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from report_content import (
    ADR_TABLE,
    ANALISE_COMPARATIVA,
    CAMADA1_CUSTO_IA,
    CAMADA2_ESFORCO_HUMANO,
    CAMADA3_CONTRAFACTUAL,
    CAPA,
    DISCLAIMER,
    SEGURANCA_TABLE,
    SECAO_10_LICOES,
    SECAO_11_INTRO,
    SECAO_11_REFERENCIAS,
    SECAO_12_APENDICES,
    SECAO_1_INTRODUCAO,
    SECAO_2_METODOLOGIA,
    SECAO_3_EXPOSICAO,
    SECAO_4_COMPOSICAO,
    SECAO_5_ENSAIO,
    SECAO_6_RESSONANCIA,
    SECAO_7_ECONOMICIDADE,
    SECAO_8_DISCUSSOES,
    SECAO_9_ETICA,
)

TEMPLATE = Path("IF1015 - Template – Relatório Final do Projeto.docx")
OUTPUT = Path("Relatorio_Final_Equipe3.docx")


def set_run_font(run, font_name="Montserrat"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def replace_paragraph_text(paragraph, new_text):
    paragraph.clear()
    run = paragraph.add_run(new_text)
    set_run_font(run)


def insert_paragraph_after(paragraph, text, style="Normal"):
    """Insert a new paragraph after the given one."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    run = new_paragraph.add_run(text)
    set_run_font(run)
    new_paragraph.style = style
    return new_paragraph


def insert_bullet_after(paragraph, text):
    """Insert a bulleted paragraph after the given one.

    O template não possui estilo de lista; usamos o prefixo '• ' como
    marcador visual e tentamos aplicar o estilo List Bullet quando disponível.
    """
    new_p = insert_paragraph_after(paragraph, f"• {text}")
    try:
        new_p.style = "List Bullet"
    except Exception:
        new_p.style = "Normal"
    return new_p


def find_paragraph(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def fill_table(table, rows):
    """Preenche uma tabela do template com lista de listas (rows)."""
    for i, row_data in enumerate(rows):
        if i >= len(table.rows):
            break
        for j, text in enumerate(row_data):
            if j >= len(table.rows[i].cells):
                break
            cell = table.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            set_run_font(run)


def insert_table_after(doc, paragraph, headers, rows):
    """Cria uma tabela nova e a insere logo após o parágrafo dado."""
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        set_run_font(run)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = ""
            run = row_cells[i].paragraphs[0].add_run(text)
            set_run_font(run)
    paragraph._element.addnext(table._element)
    return table


def fill_capa(doc):
    """Preenche os campos da capa."""
    for p in doc.paragraphs:
        if "<Nome do Projeto>" in p.text:
            replace_paragraph_text(p, CAPA["nome_projeto"])
        elif "<Nome da Equipe>" in p.text:
            replace_paragraph_text(p, CAPA["equipe"])
        elif "<URL — deve conter README.md, BUILD.md, diagramas e os 14 artefatos>" in p.text:
            replace_paragraph_text(p, CAPA["repo_url"])
        elif p.text.strip() == "<URL>" or "Sistema em produção" in p.text:
            if "<URL>" in p.text:
                replace_paragraph_text(p, CAPA["sistema_producao"])
        elif "<Nomes do integrante (login)>" in p.text:
            # Apenas o primeiro placeholder é preenchido com os nomes;
            # os demais são esvaziados para evitar duplicação.
            if not getattr(fill_capa, "_integrantes_preenchidos", False):
                replace_paragraph_text(p, "\n".join(CAPA["integrantes"]))
                fill_capa._integrantes_preenchidos = True
            else:
                replace_paragraph_text(p, "")
        elif "<Data da entrega>" in p.text:
            replace_paragraph_text(p, CAPA["data"].replace("Recife, ", ""))


def fill_disclaimer(doc):
    p = find_paragraph(doc, "Disclaimer")
    if p:
        insert_paragraph_after(p, DISCLAIMER)


def clear_template_guides(doc):
    """Remove parágrafos-guia do template após o 'Disclaimer'.

    Os parágrafos-guia explicam o que cada seção deve conter. Eles são
    removidos para que o conteúdo real seja inserido depois dos headings.
    A Seção 7 é preservada (não removida) porque seus placeholders são
    preenchidos in-place pelas subseções 7.1–7.5.
    """
    found = False
    in_section_7 = False
    to_remove = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "Disclaimer" in p.text:
            found = True
            continue
        if not found:
            continue
        if p.style.name == "Heading 2" and "7." in p.text and "Economicidade" in p.text:
            in_section_7 = True
            continue
        if in_section_7 and p.style.name == "Heading 2" and "8." in p.text:
            in_section_7 = False
            continue
        if not in_section_7 and p.style.name == "normal":
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)


def fill_economicidade(doc):
    if len(doc.tables) < 3:
        raise ValueError("Template deve conter pelo menos 3 tabelas de economicidade")
    fill_table(doc.tables[0], CAMADA1_CUSTO_IA)
    fill_table(doc.tables[1], CAMADA2_ESFORCO_HUMANO)
    fill_table(doc.tables[2], CAMADA3_CONTRAFACTUAL)


def insert_after_heading(doc, heading_needle, paragraphs):
    """Localiza um heading e insere os parágrafos logo após ele."""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and heading_needle in p.text:
            last = p
            for text in paragraphs:
                last = insert_paragraph_after(last, text)
            return last
    print(f"Aviso: heading não encontrado: {heading_needle}")
    return None


def add_subsection(paragraph, title, content):
    """Insere um subtítulo em negrito e parágrafos de conteúdo."""
    last = insert_paragraph_after(paragraph, title)
    if last.runs:
        run = last.runs[0]
        run.bold = True
        set_run_font(run)
    for text in content:
        last = insert_paragraph_after(last, text)
    return last


def fill_section_1(doc):
    insert_after_heading(doc, "1. Introdução", SECAO_1_INTRODUCAO)


def fill_section_2(doc):
    insert_after_heading(doc, "2. Metodologia", SECAO_2_METODOLOGIA)


def fill_section_3(doc):
    last = insert_after_heading(doc, "3. Movimento 1", [SECAO_3_EXPOSICAO["intro"]])
    if not last:
        return
    last = add_subsection(last, "Canvas de Estratégia e Ação (deduzido)", SECAO_3_EXPOSICAO["canvas_estrategia"])
    last = add_subsection(last, "Personas", SECAO_3_EXPOSICAO["personas"])
    last = add_subsection(last, "Missão e Visão", SECAO_3_EXPOSICAO["missao"])
    last = add_subsection(last, "Métricas de Sucesso (deduzido)", [SECAO_3_EXPOSICAO["metricas"]])
    last = add_subsection(last, "Matriz de Impacto × Esforço (deduzida)", [SECAO_3_EXPOSICAO["matriz"]])
    last = add_subsection(last, "Escopo do MVP", SECAO_3_EXPOSICAO["escopo"])


def fill_section_4(doc):
    last = insert_after_heading(doc, "4. Movimento 2", [SECAO_4_COMPOSICAO["intro"]])
    if not last:
        return
    last = add_subsection(last, "C4 Model (Níveis 1, 2 e 3)", SECAO_4_COMPOSICAO["c4"])
    last = add_subsection(last, "Registro de Decisões Arquiteturais", SECAO_4_COMPOSICAO["adrs"])

    # Tabela resumida dos ADRs
    last = insert_paragraph_after(last, "Resumo das decisões:")
    last = insert_table_after(doc, last, ADR_TABLE["headers"], ADR_TABLE["rows"])

    last = add_subsection(last, "Catálogo de Registros de Prompt", [SECAO_4_COMPOSICAO["catalogo"]])
    last = add_subsection(last, "Canvas de Design de Experimento", [SECAO_4_COMPOSICAO["experimento"]])
    last = add_subsection(last, "Protótipos, wireframes ou mockups", [SECAO_4_COMPOSICAO["prototipos"]])


def fill_section_5(doc):
    last = insert_after_heading(doc, "5. Movimento 3", [SECAO_5_ENSAIO["intro"]])
    if not last:
        return
    last = add_subsection(last, "Estratégia de desenvolvimento e tecnologias", SECAO_5_ENSAIO["dev"])
    last = add_subsection(last, "Fluxo de integração com LLMs", SECAO_5_ENSAIO["llm"])
    last = add_subsection(last, "Canvas de Testes e Validação (deduzido)", [SECAO_5_ENSAIO["testes"]])
    last = add_subsection(last, "Evidências de versionamento", [SECAO_5_ENSAIO["versionamento"]])
    last = add_subsection(last, "Análise de segurança (Aula 30)", SECAO_5_ENSAIO["seguranca"])

    # Tabela de achados de segurança
    last = insert_paragraph_after(last, "Priorização das ações:")
    last = insert_table_after(doc, last, SEGURANCA_TABLE["headers"], SEGURANCA_TABLE["rows"])

    last = add_subsection(last, "Checklist de Lançamento (deduzido)", [SECAO_5_ENSAIO["checklist"]])
    last = add_subsection(last, "Evidências de funcionamento", [SECAO_5_ENSAIO["evidencias"]])


def fill_section_6(doc):
    last = insert_after_heading(doc, "6. Movimento 4", [SECAO_6_RESSONANCIA["intro"]])
    if not last:
        return
    last = add_subsection(last, "Lançamento e coleta de feedback (deduzido)", [SECAO_6_RESSONANCIA["lancamento"]])
    last = add_subsection(last, "Painel de Feedback e Insights (deduzido)", [SECAO_6_RESSONANCIA["painel"]])
    last = add_subsection(last, "Validação das hipóteses", [SECAO_6_RESSONANCIA["hipoteses"]])
    last = add_subsection(last, "Decisão estratégica", [SECAO_6_RESSONANCIA["decisao"]])
    last = add_subsection(last, "Canvas de Escalabilidade (deduzido)", [SECAO_6_RESSONANCIA["escalabilidade"]])
    last = add_subsection(last, "Interlídio — evolução para v0.2.0", SECAO_6_RESSONANCIA["interludio"])



def _replace_after_heading(doc, heading_text, new_texts):
    """Localiza um heading e substitui o primeiro parágrafo não-vazio após ele."""
    for i, p in enumerate(doc.paragraphs):
        if heading_text in p.text and p.style.name.startswith("Heading"):
            j = i + 1
            while j < len(doc.paragraphs) and not doc.paragraphs[j].text.strip():
                j += 1
            if j < len(doc.paragraphs):
                replace_paragraph_text(doc.paragraphs[j], new_texts[0])
                last = doc.paragraphs[j]
                for text in new_texts[1:]:
                    last = insert_paragraph_after(last, text)
                return True
    return False


def fill_section_7(doc):
    # Substitui o parágrafo-guia introdutório da seção pela nota de moeda
    _replace_after_heading(doc, "7. Economicidade", [SECAO_7_ECONOMICIDADE["nota_moeda"]])

    # 7.1 — usa o heading já existente no template
    _replace_after_heading(doc, "7.1 Camada 1", [SECAO_7_ECONOMICIDADE["camada1"]])
    # A tabela do template já foi preenchida em fill_economicidade.

    # 7.2
    _replace_after_heading(doc, "7.2 Camada 2", [SECAO_7_ECONOMICIDADE["camada2"]])

    # 7.3
    _replace_after_heading(doc, "7.3 Camada 3", [SECAO_7_ECONOMICIDADE["camada3"]])

    # 7.4 — substitui os 5 parágrafos de placeholder
    comp_texts = [
        f"Custo total com IA (R$): R$ {ANALISE_COMPARATIVA['custo_com_ia_total_brl']} (tokens: R$ {ANALISE_COMPARATIVA['custo_com_ia_tokens_brl']} + horas humanas: R$ {ANALISE_COMPARATIVA['custo_com_ia_horas_brl']})",
        f"Custo total estimado sem IA (R$): R$ {ANALISE_COMPARATIVA['custo_sem_ia_brl']}",
        f"Razão de economicidade: {ANALISE_COMPARATIVA['razao']}",
        f"Saving estimado (R$): R$ {ANALISE_COMPARATIVA['saving_reais']}",
        f"Saving estimado (%): {ANALISE_COMPARATIVA['saving_percentual']}",
    ]
    idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Custo total com IA (R$)" in p.text:
            idx = i
            break
    if idx is not None:
        for k, txt in enumerate(comp_texts):
            if idx + k < len(doc.paragraphs):
                replace_paragraph_text(doc.paragraphs[idx + k], txt)

    # 7.5
    _replace_after_heading(doc, "7.5 Limitações", SECAO_7_ECONOMICIDADE["limitacoes"])


def fill_section_8(doc):
    insert_after_heading(doc, "8. Discussões Técnicas e Estratégicas", SECAO_8_DISCUSSOES)


def fill_section_9(doc):
    insert_after_heading(doc, "9. Considerações Éticas", SECAO_9_ETICA)


def fill_section_10(doc):
    last = insert_after_heading(doc, "10. Lições Aprendidas e Reflexões Finais", [SECAO_10_LICOES["sinfonia"]])
    if not last:
        return
    last = add_subsection(last, "Avaliação da proposta de valor entregue", [SECAO_10_LICOES["valor"]])
    last = add_subsection(last, "Pontos de melhoria", [SECAO_10_LICOES["melhorias"]])
    last = add_subsection(last, "Aprendizados sobre o uso de IA generativa", [SECAO_10_LICOES["ia"]])
    last = add_subsection(last, "Relato individual de cada integrante", SECAO_10_LICOES["relatos"])


def fill_section_11(doc):
    last = insert_after_heading(doc, "11. Referências", [SECAO_11_INTRO])
    for ref in SECAO_11_REFERENCIAS:
        last = insert_bullet_after(last, ref)


def fill_section_12(doc):
    last = insert_after_heading(doc, "12. Apêndices", ["Os apêndices reúnem os artefatos produzidos ao longo do projeto."])
    for item in SECAO_12_APENDICES:
        last = insert_bullet_after(last, item)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template não encontrado: {TEMPLATE}")

    shutil.copy(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)

    # 1. Limpa os parágrafos-guia do template (exceto headings).
    clear_template_guides(doc)

    # 2. Preenche capa, disclaimer e tabelas.
    fill_capa(doc)
    fill_disclaimer(doc)
    fill_economicidade(doc)

    # 3. Preenche cada seção inserindo conteúdo após o heading.
    fill_section_1(doc)
    fill_section_2(doc)
    fill_section_3(doc)
    fill_section_4(doc)
    fill_section_5(doc)
    fill_section_6(doc)
    fill_section_7(doc)
    fill_section_8(doc)
    fill_section_9(doc)
    fill_section_10(doc)
    fill_section_11(doc)
    fill_section_12(doc)

    doc.save(OUTPUT)
    print(f"Relatório gerado: {OUTPUT}")


if __name__ == "__main__":
    main()
